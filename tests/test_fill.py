import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "examples"))

from create_sample_template import build_sample_template  # noqa: E402
from masterform import fill_document, fill_template, find_placeholders  # noqa: E402

DATA = {
    "anrede": "Herr",
    "name": "Max Mustermann",
    "firma": "ACME GmbH",
    "eintrittsdatum": "01.01.2024",
    "unterzeichner": "Die Geschäftsführung",
    "personalnummer": "12345",
    "datum": "13.08.2026",
}


def _all_text(doc: Document) -> str:
    parts = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(cell.paragraphs)
    for section in doc.sections:
        parts.extend(section.footer.paragraphs)
    return "\n".join(p.text for p in parts)


def test_find_placeholders(tmp_path):
    template_path = tmp_path / "vorlage.docx"
    build_sample_template().save(template_path)

    placeholders = find_placeholders(str(template_path))

    assert placeholders == set(DATA.keys())


def test_fill_template_replaces_all_known_placeholders(tmp_path):
    template_path = tmp_path / "vorlage.docx"
    output_path = tmp_path / "ausgefuellt.docx"
    build_sample_template().save(template_path)

    fill_template(str(template_path), str(output_path), DATA)

    result = Document(str(output_path))
    text = _all_text(result)

    assert "{{" not in text
    assert "Herr Max Mustermann" in text
    assert "ACME GmbH" in text
    assert "12345" in text
    assert "13.08.2026" in text


def test_unknown_placeholder_stays_untouched(tmp_path):
    template_path = tmp_path / "vorlage.docx"
    output_path = tmp_path / "ausgefuellt.docx"
    build_sample_template().save(template_path)

    partial_data = {k: v for k, v in DATA.items() if k != "personalnummer"}
    fill_template(str(template_path), str(output_path), partial_data)

    text = _all_text(Document(str(output_path)))
    assert "{{personalnummer}}" in text


def test_run_split_placeholder_is_still_found():
    """Simuliert einen Platzhalter, der über mehrere Runs verteilt ist."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Hallo {{na")
    p.add_run("me}}, willkommen!")

    fill_document(doc, {"name": "Anna"})

    assert p.text == "Hallo Anna, willkommen!"

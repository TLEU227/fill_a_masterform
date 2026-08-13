#!/usr/bin/env python3
"""Erzeugt eine Beispiel-Textvorlage (examples/vorlage.docx) zum Ausprobieren.

Word-Dateien sind Binärdateien und lassen sich nicht sinnvoll als reiner
Text im Repo pflegen - daher wird die Beispielvorlage per Skript erzeugt.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

OUTPUT_PATH = Path(__file__).parent / "vorlage.docx"


def build_sample_template() -> Document:
    doc = Document()

    doc.add_heading("Bestätigungsschreiben", level=1)

    p = doc.add_paragraph()
    p.add_run("Sehr geehrte/r ").bold = False
    p.add_run("{{anrede}} {{name}}").bold = True
    p.add_run(",")

    doc.add_paragraph(
        "hiermit bestätigen wir, dass {{name}} bei {{firma}} seit dem "
        "{{eintrittsdatum}} beschäftigt ist."
    )
    doc.add_paragraph("Mit freundlichen Grüßen")
    doc.add_paragraph("{{unterzeichner}}")

    table = doc.add_table(rows=2, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Feld"
    hdr[1].text = "Wert"
    row = table.rows[1].cells
    row[0].text = "Personalnummer"
    row[1].text = "{{personalnummer}}"

    section = doc.sections[0]
    section.footer.paragraphs[0].text = "Erstellt am {{datum}} - {{firma}}"

    return doc


if __name__ == "__main__":
    build_sample_template().save(OUTPUT_PATH)
    print(f"Beispielvorlage erstellt: {OUTPUT_PATH}")

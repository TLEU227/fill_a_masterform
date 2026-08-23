"""Demo v8: wie v7, PLUS durchgängige Nummerierung + echte Ankreuzfelder.

NEU (Nutzer-Wunsch):
1. Jede automatisch eingesetzte/ausgewählte Stelle bekommt eine eindeutige
   Nummer nach dem Muster "<Ort>-<laufende Nr.>", z.B. "Tabelle 1-1",
   "Kap. 1.1-2" - direkt sichtbar im Dokumenttext (in eckigen Klammern vor
   dem Wert). Zusätzlich wird eine Zuordnungstabelle (Nr. -> DB-Feld/Ort ->
   Wert) erzeugt, siehe ZUORDNUNGSDATEI am Ende.
2. Tabelle 10 (Weitere Validierungsdokumente) hat als einzige Stelle im
   CS-VP echte Word-Ankreuzfelder (w14:checkbox Content Controls, 18 Zeilen
   x Ja/Nein = 36 Stück - NICHT in Tabellen 1-9/11/12 vorhanden, das sind
   alles normale Absatz-/Zellentexte). Angekreuzte Kästchen werden gelb
   markiert (☐->☑ + Highlight), analog zur bisherigen Farbkonvention.
   NEU ENTDECKT: nur 1 der 18 Zeilen (PPQ) hat aktuell ein passendes
   Projekt-DB-Feld (phase_ppq_geplant) - die anderen 17 (Datenflussdiagramm,
   Audit Trail Review Konzept, Berechtigungskonzept, Trainingsplan, UPM,
   Datenmigration, Wartung/Monitoring, Archivierung, Backup&Restore,
   Business Continuity, Incident-Management, Change-/Konfigurations-
   management, Logbuch, Lieferantenbewertung, Quality Agreement,
   Bedienungsanweisungen, Handbuch) haben noch KEIN Feld in der Projekt-DB -
   werden deshalb bewusst nicht angekreuzt (keine erfundenen GxP-Aussagen),
   sondern nur in "skipped" vermerkt.

Farbcode wie zuvor: gelb = automatisch eingefügt/geändert/angekreuzt,
durchgestrichen + dunkelrot = zum Löschen vorgeschlagen (nichts wird real
entfernt, siehe mark_as_deleted()). Die "[Tabelle X-N]"-Nummerierung ist
eine reine Demo-/Abstimmungshilfe (blau, klein, kursiv) - kein Teil des
eigentlichen Dokumenttexts, vor echter Nutzung wieder zu entfernen.
"""
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
import re

SRC = "/root/.claude/uploads/4bad0ed6-f59f-58ad-8101-0f3d1b25ac47/69d6102d-FFICFQUMT0000722MDT__CSVP_Validierungsplan_9.0.docx"
OUT = "/tmp/claude-0/-home-user-fill-a-masterform/4bad0ed6-f59f-58ad-8101-0f3d1b25ac47/scratchpad/demo/CS-VP_AUSGEFUELLT_DEMO_v8.docx"
ZUORDNUNGSDATEI = "/tmp/claude-0/-home-user-fill-a-masterform/4bad0ed6-f59f-58ad-8101-0f3d1b25ac47/scratchpad/demo/CS-VP_Zuordnungstabelle_demo8.md"

# ============================================================ System-DB (MLCS 1428, wie bisher)
SYS = {
    "gebaeude": "G680", "bereich": "MIB Produktion 2",
    "systemname": "SPS Zentrifugentrockner PU4300A", "mlcs_id": "1428",
    "kurzbeschreibung": "Speicherprogrammierbare Steuerung (SPS) für PU Zentrifugentrockner 4330A",
    "sw_name": "Simatic", "sw_version": "SV=VV: V3.0.2 / Speicherprogrammierbare Steuerung (SPS)",
    "hersteller": "Fima",
    "gxp_kritikalitaet": "Minor", "gamp_kategorie": "4", "eres_typ_nr": "4",
    "systemtyp": "CE-PCS",
    "lieferantennummer": "OOZ000000062125",
    "ki_reifegrad": "N/A",
    "testtiefe": "Gering",
}

# ============================================================ Personen (System-DB, Objektart 'person')
PERSONEN = {
    "rolle_ersteller": {"name": "Julia Fischer", "funktion": "C&Q Engineer"},
    "rolle_sme": {"name": "Michael Bauer", "funktion": "SME Automatisierungstechnik"},
    "rolle_si_pl": {"name": "Anna Keller", "funktion": "Projektleiterin Retrofit PU4300A"},
    "rolle_tso": {"name": "Thomas Weber", "funktion": "Technical System Owner PU43xx"},
    "rolle_bso": {"name": "Sabine Hoffmann", "funktion": "Business System Owner MIB Produktion 2"},
    "rolle_bqr": {"name": "Peter Schmidt", "funktion": "Quality Representative CSV"},
    # rolle_csq bewusst NICHT belegt.
}

# ============================================================ Projekt-DB (fiktiver Demo-Datensatz)
PROJ = {
    "projektbezeichnung": "Retrofit PU4300A",
    "ist_folgeprojekt": "ja",
    "vorgaenger_dok_id": "FRA-PLAN-G-00421", "vorgaenger_version": "2.0",
    "folgeversion": "3.0",
    "change_control_nummer": "CC-2024-005123",
    "hauptfunktion": "Steuerung und Überwachung des Trocknungsprozesses am Zentrifugentrockner PU4300A, inkl. Rezeptverwaltung und Prozessdatenerfassung.",
    "systembewertung_dok_id": "FRA-PLAN-G-00987", "systembewertung_version": "2.0",
    "urs_dok_id": "FRA-PLAN-G-00654", "urs_version": "2.0",
    "fs_dok_id": "FRA-PLAN-G-00655", "fs_version": "2.0",
    "vmp_erforderlich": "nein",
    "digital_beteiligt": "nein",
    "urs_bereits_erstellt": "ja",
    "phasen_geplant": {"DQ": True, "IQ": True, "OQ": True, "PQ": True, "PPQ": False},
    "gep_pruefung_erforderlich": "nein",
    "testplan_art": "als_anhang",
    "testdurchfuehrung_art": "lieferant_durchfuehrung_sanofi_aufsicht",
    "verantwortlich_funktionsspezifikation": "Lieferant",
    "verantwortlich_hds": "Sanofi",
    "verantwortlich_sds": "Lieferant",
    "verantwortlich_ra": "Sanofi",
    "verantwortlich_tm": "Sanofi",
    "verantwortlich_iq_testvorschriften": "Lieferant",
    "verantwortlich_iq_durchfuehrung": "Lieferant",
}
LIEFERANT_VERANTWORTLICHKEITEN = [
    "die Erstellung der Spezifikationen",
    "die technische Umsetzung der Anforderungen",
    "die Durchführung der Validierung",
    "die Schulung des Sanofi-Personals am System",
    "die Erstellung der Anwenderdokumentation",
]
VERSIONSHISTORIE = [
    {"version": "2.0", "cc_nummer": "CC-2022-000456", "beschreibung": "Anpassung der Steuerungssoftware nach Austausch der SPS-Hardware."},
    {"version": "3.0", "cc_nummer": "CC-2024-005123", "beschreibung": "Retrofit des Antriebssystems und Anpassung der Rezeptparameter."},
]

changes = []
skipped = []

# ============================================================ Nummerierung (NEU in v8)
# "<Ort>-<laufende Nr>" je Bereich, direkt im Dokument sichtbar (kleine,
# kursive, blaue Markierung "[Ort-Nr]" vor dem eingesetzten Wert bzw. am
# Ende eines beibehaltenen Absatzes/einer Checkbox-Zeile). Reine Demo-/
# Abstimmungshilfe, siehe ZUORDNUNGSDATEI fuer die vollstaendige Legende.
NR_COUNTER = {}
NUMMERIERUNG = []  # (tag, label, wert) fuer die Zuordnungstabelle

def nr(bereich):
    NR_COUNTER[bereich] = NR_COUNTER.get(bereich, 0) + 1
    return f"{bereich}-{NR_COUNTER[bereich]}"

def merke_nummer(tag, label, wert):
    NUMMERIERUNG.append((tag, label, wert))

# ============================================================ Helper (aus v7 uebernommen + bereich-Parameter)
DELETE_COLOR = RGBColor(0xC0, 0x00, 0x00)
NR_TAG_COLOR = RGBColor(0x00, 0x33, 0x99)

def mark_as_deleted(paragraph):
    for r in paragraph.runs:
        if not r.text:
            continue
        r.font.strike = True
        r.font.color.rgb = DELETE_COLOR
        r.font.highlight_color = None

def add_tag_run(paragraph, tag):
    """Haengt ' [Ort-Nr]' als kleinen, blauen, kursiven Referenz-Tag ans Ende
    eines Absatzes an - reine Zuordnungshilfe fuer dieses Demo, kein Teil des
    eigentlichen Dokumenttexts."""
    run = paragraph.add_run(f" [{tag}]")
    run.font.italic = True
    run.font.size = Pt(7)
    run.font.color.rgb = NR_TAG_COLOR

def mark_as_kept(paragraph, bereich=None, skip_grey=True):
    """Markiert einen Absatz als 'trifft zu' (schwarz+gelb). Graue Runs
    (A6A6A6) werden NICHT angefasst. Mit bereich=... wird zusaetzlich eine
    Nummer vergeben und als Tag ans Ende angehaengt."""
    for r in paragraph.runs:
        if not r.text:
            continue
        if skip_grey and r.font.color and r.font.color.type and r.font.color.rgb == RGBColor(0xA6, 0xA6, 0xA6):
            continue
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if bereich:
        tag = nr(bereich)
        add_tag_run(paragraph, tag)
        return tag
    return None

def delete_paragraph(paragraph, label):
    old = paragraph.text
    mark_as_deleted(paragraph)
    changes.append((label, old, "[durchgestrichen/rot markiert - zum Löschen vorgeschlagen, nicht real entfernt]"))

def replace_span_in_paragraph(paragraph, start, end, value, mark_black=False):
    runs = paragraph.runs
    offsets, pos = [], 0
    for r in runs:
        offsets.append((pos, pos + len(r.text)))
        pos += len(r.text)
    overlapping = [i for i, (s, e) in enumerate(offsets) if e > start and s < end]
    if not overlapping:
        return False
    first_i, last_i = overlapping[0], overlapping[-1]
    first_run = runs[first_i]
    prefix = first_run.text[: start - offsets[first_i][0]]
    if first_i == last_i:
        suffix = first_run.text[end - offsets[first_i][0]:]
        first_run.text = prefix + value + suffix
    else:
        last_run = runs[last_i]
        suffix = last_run.text[end - offsets[last_i][0]:]
        first_run.text = prefix + value
        for i in overlapping[1:-1]:
            runs[i].text = ""
        last_run.text = suffix
    if mark_black:
        first_run.font.color.rgb = RGBColor(0, 0, 0)
        first_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        first_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        first_run.font.color.rgb = None
    return True

def replace_marker(paragraph, marker, value, label, mark_black=False, bereich=None):
    text = paragraph.text
    idx = text.find(marker)
    if idx == -1:
        return None
    tag = nr(bereich) if bereich else None
    display_value = f"[{tag}] {value}" if tag else value
    ok = replace_span_in_paragraph(paragraph, idx, idx + len(marker), display_value, mark_black=mark_black)
    if ok:
        changes.append((label, marker, display_value))
        if tag:
            merke_nummer(tag, label, value)
    return tag if ok else None

def replace_whole_paragraph(paragraph, value, label, bereich=None):
    old = paragraph.text
    tag = nr(bereich) if bereich else None
    display_value = f"[{tag}] {value}" if tag else value
    for r in paragraph.runs[1:]:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = display_value
        paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        paragraph.runs[0].font.color.rgb = None
        changes.append((label, old, display_value))
        if tag:
            merke_nummer(tag, label, value)
        return tag
    return None

def set_cell_value(cell, value, label, mark_black=False, bereich=None):
    old = cell.text
    tag = nr(bereich) if bereich else None
    display_value = f"[{tag}] {value}" if tag else value
    runs = [r for p in cell.paragraphs for r in p.runs]
    if not runs:
        run = cell.paragraphs[0].add_run(display_value)
    else:
        run = runs[0]
        run.text = display_value
        for r in runs[1:]:
            r.text = ""
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.color.rgb = RGBColor(0, 0, 0) if mark_black else None
    changes.append((label, old.strip(), display_value))
    if tag:
        merke_nummer(tag, label, value)
    return tag

def append_to_cell(cell, value, label, mark_black=True, bereich=None):
    old = cell.text
    tag = nr(bereich) if bereich else None
    display_value = f"[{tag}] {value}" if tag else value
    p = cell.paragraphs[-1]
    run = p.add_run(" " + display_value)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if mark_black:
        run.font.color.rgb = RGBColor(0, 0, 0)
    changes.append((label, old.strip(), (old.strip() + " " + display_value).strip()))
    if tag:
        merke_nummer(tag, label, value)
    return tag

def find_p(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p):
            return p
    return None

def find_all_p(doc, predicate):
    return [p for p in doc.paragraphs if predicate(p)]

def delete_block(paragraphs, label):
    for p in paragraphs:
        if p.text.strip():
            mark_as_deleted(p)
    if paragraphs:
        changes.append((label, paragraphs[0].text[:60], "[Block durchgestrichen/rot - zum Löschen vorgeschlagen]"))

def block_between(doc, start_pred, end_is_heading=True):
    paras = doc.paragraphs
    start_i = next(i for i, p in enumerate(paras) if start_pred(p))
    end_i = start_i + 1
    while end_i < len(paras) and not paras[end_i].style.name.startswith("Heading"):
        end_i += 1
    return paras[start_i:end_i]

# ============================================================ NEU in v8: Ankreuzfelder (w14:checkbox Content Controls)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"

def _sdts_in_cell(cell):
    return cell._tc.findall(f".//{{{W_NS}}}sdt")

def _set_checkbox(sdt_element, checked):
    checked_el = sdt_element.find(f".//{{{W14_NS}}}checked")
    checked_el.set(f"{{{W14_NS}}}val", "1" if checked else "0")
    t_el = sdt_element.find(f".//{{{W_NS}}}t")
    t_el.text = "☒" if checked else "☐"
    if checked:
        r_el = t_el.getparent()
        rpr = r_el.find(f"{{{W_NS}}}rPr")
        if rpr is None:
            rpr = r_el.makeelement(f"{{{W_NS}}}rPr", {})
            r_el.insert(0, rpr)
        hl = rpr.find(f"{{{W_NS}}}highlight")
        if hl is None:
            hl = rpr.makeelement(f"{{{W_NS}}}highlight", {})
            rpr.append(hl)
        hl.set(f"{{{W_NS}}}val", "yellow")

def set_ja_nein_checkbox(cell, antwort, label, bereich):
    """cell = die Tabellenzelle mit BEIDEN Checkboxen ('☐ Ja / ☐ Nein' als
    zwei w:sdt in EINEM Zellen-Absatz). antwort: 'ja' oder 'nein' - genau
    diese wird angekreuzt+gelb markiert, die andere bleibt unangetastet
    (leeres Kaestchen). Haengt zusaetzlich den Nummer-Tag ans Zellenende."""
    sdts = _sdts_in_cell(cell)
    assert len(sdts) == 2, f"Erwartet genau 2 Checkboxen in der Zelle, gefunden: {len(sdts)}"
    ja_sdt, nein_sdt = sdts
    _set_checkbox(ja_sdt, antwort == "ja")
    _set_checkbox(nein_sdt, antwort == "nein")
    tag = nr(bereich)
    add_tag_run(cell.paragraphs[-1], tag)
    changes.append((label, "☐ Ja / ☐ Nein (nicht angekreuzt)", f"[{tag}] {'☑ Ja' if antwort == 'ja' else '☑ Nein'} (angekreuzt + gelb markiert)"))
    merke_nummer(tag, label, antwort.capitalize())
    return tag

doc = Document(SRC)

# ============================================================ 1. Kopf-Tabelle
for t in doc.tables:
    if t.rows[0].cells[0].text.strip().startswith("Gebäude"):
        set_cell_value(t.rows[0].cells[1], SYS["gebaeude"], "Kopf-Tabelle: Gebäude", bereich="Kopf-Tabelle")
        set_cell_value(t.rows[0].cells[3], SYS["bereich"], "Kopf-Tabelle: Bereich", bereich="Kopf-Tabelle")
        set_cell_value(t.rows[1].cells[1], SYS["systemname"], "Kopf-Tabelle: Systemname", bereich="Kopf-Tabelle")
        set_cell_value(t.rows[1].cells[3], SYS["mlcs_id"], "Kopf-Tabelle: MLCS-ID", bereich="Kopf-Tabelle")
        break

# ============================================================ 2. Tabelle 1: Dokumentenfreigabe - Personen ergaenzen
def person_suffix(rollen_keys):
    teile = []
    for key in rollen_keys:
        p = PERSONEN.get(key)
        if p:
            teile.append(f"{p['name']}, {p['funktion']}")
    if not teile:
        return None
    return "— " + "; ".join(teile)

TABELLE1_ZEILEN = [
    ("Autor", None, ["rolle_ersteller"]),
    ("Prüfer", "Projektleiter/SME", ["rolle_si_pl", "rolle_sme"]),
    ("Prüfer", "Technical System Owner (TSO)", ["rolle_tso"]),
    ("Prüfer", "Business System Owner (BSO)", ["rolle_bso"]),
    ("Freigeber", "Business Quality Representative (BQR)", ["rolle_bqr"]),
    ("Freigeber", "Computerized System Quality Expert (CSQ)", ["rolle_csq"]),
]
for t in doc.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if header == ["Rolle", "Funktion", "Beschreibung"]:
        for row in t.rows[1:]:
            rolle_zelle = row.cells[0].text.strip()
            funktion_zelle = row.cells[1].text.strip()
            eintrag = next(
                (e for e in TABELLE1_ZEILEN if e[0] == rolle_zelle and (e[1] is None or e[1] == funktion_zelle)),
                None,
            )
            if eintrag is None:
                continue
            _, _, rollen_keys = eintrag
            suffix = person_suffix(rollen_keys)
            if suffix:
                append_to_cell(row.cells[1], suffix, f"Tabelle 1: Dokumentenfreigabe - Funktion-Spalte '{funktion_zelle}' um Person(en) ergänzt ({', '.join(rollen_keys)}, System-DB)", bereich="Tabelle 1")
            else:
                skipped.append((f"Tabelle 1: Dokumentenfreigabe - Zeile '{rolle_zelle} / {funktion_zelle}'", f"Keine Person für {', '.join(rollen_keys)} zugewiesen -> Zeile bleibt unverändert."))
        break

# ============================================================ 3. Kap. 1.1 Ziel und Umfang
p = find_p(doc, lambda p: "<<System, MLCS ID>>" in p.text)
if p:
    replace_marker(p, "<<System, MLCS ID>>", f"{SYS['systemname']}, MLCS-ID {SYS['mlcs_id']}", "Kap. 1.1: Systemname/MLCS-ID im Fließtext", bereich="Kap. 1.1")

p = find_p(doc, lambda p: p.text.strip() == "System zur / zum....")
if p:
    replace_whole_paragraph(p, SYS["kurzbeschreibung"], "Kap. 1.1: Systembeschreibung (Kurzbeschreibung, System-DB)", bereich="Kap. 1.1")

p = find_p(doc, lambda p: "Change Controls XXXXXXX" in p.text)
if p:
    replace_marker(p, "XXXXXXX", PROJ["change_control_nummer"], "Kap. 1.1: Change-Control-Nummer dieses Projekts (Projekt-DB)", bereich="Kap. 1.1")

# ============================================================ 4. Kap. 1.4 Systembeschreibung
p = find_p(doc, lambda p: p.text.strip().startswith("Das System dient der"))
if p:
    replace_marker(p, "…(Zusammenfassende Beschreibung der Haupt-Funktion)", PROJ["hauptfunktion"], "Kap. 1.4: Hauptfunktion des Systems (projektspezifisch, Projekt-DB) - Einleitung 'Das System dient der' bleibt erhalten", bereich="Kap. 1.4")

p = find_p(doc, lambda p: "wird im Gebäude" in p.text)
if p:
    replace_marker(p, "XXX", SYS["gebaeude"], "Kap. 1.4: Betriebsort - Gebäude", bereich="Kap. 1.4")
    skipped.append(("Kap. 1.4: Betriebsort - Raum", "Feld raum ist für MLCS 1428 leer -> 'Raum XXX' bleibt unangetastet."))

# ============================================================ 5. Kap. 1.5 Systembewertung (Projekt-DB!)
p = find_p(doc, lambda p: "durchgeführt (XXXX V x.y)" in p.text)
if p:
    replace_marker(p, "XXXX V x.y", f"{PROJ['systembewertung_dok_id']} V {PROJ['systembewertung_version']}", "Kap. 1.5: Systembewertung Dok-ID + Version (Projekt-DB, nicht System-DB!)", bereich="Kap. 1.5")

p = find_p(doc, lambda p: p.text.strip() == "MLCS-ID:")
if p:
    tag = nr("Kap. 1.5")
    run = p.add_run(f"[{tag}] {SYS['mlcs_id']}")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if p.runs:
        run.font.name = p.runs[0].font.name
        run.font.size = p.runs[0].font.size
    changes.append(("Kap. 1.5: Systemeinstufung - MLCS-ID", "MLCS-ID: ", f"MLCS-ID: [{tag}] " + SYS["mlcs_id"]))
    merke_nummer(tag, "Kap. 1.5: Systemeinstufung - MLCS-ID", SYS["mlcs_id"])

p = find_p(doc, lambda p: "GMP-Kritikalität" in p.text and "Minor, Major, Critical" in p.text)
if p:
    replace_marker(p, "Minor, Major, Critical", SYS["gxp_kritikalitaet"], "Kap. 1.5: Systemeinstufung - GMP-Kritikalität", bereich="Kap. 1.5")

p = find_p(doc, lambda p: p.text.strip().startswith("Softwarekategorie: X"))
if p:
    replace_marker(p, "Softwarekategorie: X", f"Softwarekategorie: {SYS['gamp_kategorie']}", "Kap. 1.5: Systemeinstufung - Softwarekategorie (GAMP 5)", bereich="Kap. 1.5")

p = find_p(doc, lambda p: p.text.strip().startswith("eRecord & eSignature Typ: X"))
if p:
    replace_marker(p, "Typ: X", f"Typ: {SYS['eres_typ_nr']}", "Kap. 1.5: Systemeinstufung - eRecord/eSignature-Typ (ERES)", bereich="Kap. 1.5")

p = find_p(doc, lambda p: p.text.strip().startswith("CS-Kategorisierung: X"))
if p:
    replace_marker(p, "CS-Kategorisierung: X", f"CS-Kategorisierung: {SYS['systemtyp']}", "Kap. 1.5: Systemeinstufung - CS-Kategorisierung (Systemtyp)", bereich="Kap. 1.5")

TESTTIEFE_LEGENDE = {
    "Gering": "Gering: Exploratory Testing",
    "Mittel": "Mittel: Exploratory und Nominal Testing",
    "Hoch": "Hoch: Exploratory, Nominal und Challenge Testing",
}
zutreffende_zeile = TESTTIEFE_LEGENDE[SYS["testtiefe"]]
for p in find_all_p(doc, lambda p: p.text.strip() in TESTTIEFE_LEGENDE.values() and p.text.strip() != zutreffende_zeile):
    delete_paragraph(p, f"Kap. 1.5: Testtiefe-Legende (berechnete Testtiefe = {SYS['testtiefe']})")
p = find_p(doc, lambda p: p.text.strip() == zutreffende_zeile)
if p:
    tag = mark_as_kept(p, bereich="Kap. 1.5")
    changes.append((f"Kap. 1.5: Testtiefe-Legende - zutreffende Zeile (Testtiefe={SYS['testtiefe']})", zutreffende_zeile, zutreffende_zeile))
    merke_nummer(tag, "Kap. 1.5: Testtiefe-Legende - zutreffende Zeile", zutreffende_zeile)

# ============================================================ 6. Kap. 1.6 Systemgrenzen
p = find_p(doc, lambda p: "im folgenden Dokument XXXX festgelegt" in p.text)
if p:
    replace_marker(p, "XXXX", PROJ["systembewertung_dok_id"], "Kap. 1.6: Systembewertung Dok-ID (Systemgrenzen-Referenz, Projekt-DB)", bereich="Kap. 1.6")

# ============================================================ 7. Kap. 1.7 Schnittstellen (Stichwortsuche)
systemtext = " ".join([SYS.get("kurzbeschreibung", ""), SYS.get("sw_name", ""), SYS.get("systemname", "")]).lower()
SCHNITTSTELLEN_REGELN = [
    (["tcp/ip", "tcp-ip"], "Über eine TCP/IP-Verbindung stellt der Client dem System Daten für die Protokollierung und Archivierung zur Verfügung."),
    (["prodis"], "Es wird geprüft, dass die Datenübertragung an das validierte PRODIS System (IP21) korrekt erfolgt."),
    (["sap"], "Es wird geprüft, dass die Datenübertragung an das validierte SAP System korrekt erfolgt."),
    (["sanofi laufwerk", "laufwerk"], "Es wird geprüft, dass die Datenübertragung auf Sanofi Laufwerke korrekt erfolgt."),
]
gefunden = 0
for keywords, zeilentext in SCHNITTSTELLEN_REGELN:
    if any(kw in systemtext for kw in keywords):
        gefunden += 1
        p = find_p(doc, lambda p, zt=zeilentext: p.text.strip() == zt)
        if p:
            tag = mark_as_kept(p, bereich="Kap. 1.7")
            changes.append(("Kap. 1.7: Schnittstellen-Aufzählung - Zeile trifft laut Stichwortsuche zu", zeilentext, zeilentext))
            merke_nummer(tag, "Kap. 1.7: Schnittstellen-Aufzählung", zeilentext[:50])
if gefunden == 0:
    skipped.append(("Kap. 1.7: Schnittstellen-Aufzählung", "Keine der Stichwörter (TCP/IP, PRODIS, SAP, Sanofi Laufwerke) kommt in den Textfeldern von MLCS 1428 vor -> alle 4 Zeilen bleiben unverändert."))

# ============================================================ 8. Tabelle 6: Mitgeltende Unterlagen
for t in doc.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if header == ["Dok-ID", "Beschreibung", "Freigabedatum", "Version"]:
        letzte_zeile = None
        for row in t.rows[1:]:
            beschr = row.cells[1].text
            if "Systembewertung gemäß" in beschr:
                set_cell_value(row.cells[0], PROJ["systembewertung_dok_id"], "Tabelle 6: Mitgeltende Unterlagen - Dok-ID Systembewertung (Projekt-DB)", bereich="Tabelle 6")
                set_cell_value(row.cells[3], PROJ["systembewertung_version"], "Tabelle 6: Mitgeltende Unterlagen - Version Systembewertung (Projekt-DB)", bereich="Tabelle 6")
            elif "Validerungs-Master-Plan" in beschr or "Validierungs-Master-Plan" in beschr:
                if PROJ["vmp_erforderlich"] == "ja" and PROJ.get("vmp_dok_id"):
                    set_cell_value(row.cells[0], PROJ["vmp_dok_id"], "Tabelle 6: Mitgeltende Unterlagen - Dok-ID VMP", bereich="Tabelle 6")
                    if PROJ.get("vmp_version"):
                        set_cell_value(row.cells[3], PROJ["vmp_version"], "Tabelle 6: Mitgeltende Unterlagen - Version VMP", bereich="Tabelle 6")
                else:
                    for p in row.cells[0].paragraphs + row.cells[1].paragraphs:
                        mark_as_deleted(p)
                    changes.append(("Tabelle 6: Zeile 'Validierungs-Master-Plan' gestrichen (vmp_erforderlich=nein bzw. keine Dok-ID, Projekt-DB)", beschr.strip(), "[durchgestrichen/rot - zum Löschen vorgeschlagen]"))
            elif "URS" in beschr:
                if PROJ.get("urs_dok_id"):
                    set_cell_value(row.cells[0], PROJ["urs_dok_id"], "Tabelle 6: Mitgeltende Unterlagen - Dok-ID URS (Projekt-DB)", bereich="Tabelle 6")
                    if PROJ.get("urs_version"):
                        set_cell_value(row.cells[3], PROJ["urs_version"], "Tabelle 6: Mitgeltende Unterlagen - Version URS (Projekt-DB)", bereich="Tabelle 6")
                else:
                    skipped.append(("Tabelle 6: Zeile 'URS'", "Kein urs_dok_id in der Projekt-DB hinterlegt -> Zeile bleibt unangetastet."))
            letzte_zeile = row
        if PROJ.get("fs_dok_id") and letzte_zeile is not None:
            new_row = t.add_row()
            set_cell_value(new_row.cells[0], PROJ["fs_dok_id"], "Tabelle 6: Mitgeltende Unterlagen - neue Zeile Funktionsspezifikation, Dok-ID (Projekt-DB)", mark_black=True, bereich="Tabelle 6")
            set_cell_value(new_row.cells[1], "Funktionsspezifikation", "Tabelle 6: Mitgeltende Unterlagen - neue Zeile Funktionsspezifikation, Beschreibung", mark_black=True, bereich="Tabelle 6")
            if PROJ.get("fs_version"):
                set_cell_value(new_row.cells[3], PROJ["fs_version"], "Tabelle 6: Mitgeltende Unterlagen - neue Zeile Funktionsspezifikation, Version (Projekt-DB)", mark_black=True, bereich="Tabelle 6")
        elif not PROJ.get("fs_dok_id"):
            skipped.append(("Tabelle 6: neue Zeile 'Funktionsspezifikation'", "Kein fs_dok_id in der Projekt-DB hinterlegt -> keine Zeile ergänzt."))
        break

# ============================================================ 9. Kap. 2.2 Verantwortlichkeiten des Lieferanten
p = find_p(doc, lambda p: "<<MUSTER>>" in p.text)
if p:
    replace_marker(p, "<<MUSTER>>", SYS["hersteller"], "Kap. 2.2: Verantwortlichkeiten des Lieferanten - Firmenname (<<MUSTER>>)", bereich="Kap. 2.2")

TEMPLATE_VERANTWORTLICHKEITEN = [
    "die Erstellung der Spezifikationen",
    "die technische Umsetzung der Anforderungen",
    "die ordnungsgemäße Installation des Systems",
    "die Durchführung der Validierung",
]
last_bullet_p = None
for text in TEMPLATE_VERANTWORTLICHKEITEN:
    p = find_p(doc, lambda p, t=text: p.text.strip() == t)
    if not p:
        continue
    if text in LIEFERANT_VERANTWORTLICHKEITEN:
        tag = mark_as_kept(p, bereich="Kap. 2.2")
        changes.append(("Kap. 2.2: Verantwortlichkeit des Lieferanten - ausgewählt (Projekt-DB)", text, text))
        merke_nummer(tag, "Kap. 2.2: Verantwortlichkeit des Lieferanten", text[:50])
    else:
        delete_paragraph(p, "Kap. 2.2: Verantwortlichkeit des Lieferanten - NICHT ausgewählt (Projekt-DB)")
    last_bullet_p = p
fehlende = [v for v in LIEFERANT_VERANTWORTLICHKEITEN if v not in TEMPLATE_VERANTWORTLICHKEITEN]
if fehlende and last_bullet_p is not None:
    anchor = last_bullet_p._element
    for value in fehlende:
        tag = nr("Kap. 2.2")
        new_p = doc.add_paragraph(style=last_bullet_p.style)
        new_p.paragraph_format.left_indent = last_bullet_p.paragraph_format.left_indent
        run = new_p.add_run(f"[{tag}] {value}")
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        anchor.addnext(new_p._element)
        anchor = new_p._element
        changes.append(("Kap. 2.2: Verantwortlichkeit des Lieferanten - NEU ergänzt (Projekt-DB, im Template nicht als Zeile vorgesehen)", "-", f"[{tag}] {value}"))
        merke_nummer(tag, "Kap. 2.2: Verantwortlichkeit des Lieferanten - NEU ergänzt", value[:50])

# ============================================================ 10. Kap. 3.1 VMP
if PROJ["vmp_erforderlich"] == "nein":
    block = block_between(doc, lambda p: p.style.name == "Heading 2" and "Validation Master Plan" in p.text)
    delete_block(block, "Kap. 3.1: Validierungsstrategie gemäß VMP - komplett gestrichen (vmp_erforderlich=nein, Projekt-DB)")
else:
    p = find_p(doc, lambda p: "Gemäß Doc ID-xxx" in p.text)
    if p and PROJ.get("vmp_dok_id"):
        replace_marker(p, "Doc ID-xxx", PROJ["vmp_dok_id"], "Kap. 3.1: VMP Dok-ID (Projekt-DB)", bereich="Kap. 3.1")

# ============================================================ 11. Kap. 3.2 Einsatz von Künstlicher Intelligenz
KI_BLOCK = {"N/A": "OHNE", "I": "OHNE", "II": "OHNE", "III": "MIT", "IV": "MIT"}
ki_wert = SYS["ki_reifegrad"]
block_art = KI_BLOCK.get(ki_wert)
if block_art is None:
    skipped.append(("Kap. 3.2: Einsatz von Künstlicher Intelligenz", f"ki_reifegrad='{ki_wert}' (V/VI) ist laut Template im GxP-Umfeld nicht zugelassen - bewusst nicht automatisch befüllt."))
else:
    paras = doc.paragraphs
    ohne_idx = next(i for i, p in enumerate(paras) if p.text.strip() == "OHNE künstliche Intelligenz")
    mit_idx = next(i for i, p in enumerate(paras) if p.text.strip() == "Oder MIT Künstliche Intelligenz:")
    end_idx = mit_idx + 1
    while not paras[end_idx].style.name.startswith("Heading"):
        end_idx += 1
    ohne_range = paras[ohne_idx:mit_idx]
    mit_range = paras[mit_idx:end_idx]
    zutreffender_block, nicht_zutreffender_block = (ohne_range, mit_range) if block_art == "OHNE" else (mit_range, ohne_range)

    ai_table = next((t for t in doc.tables if [c.text.strip() for c in t.rows[0].cells] == ["IDs", "AI Standard", "AI High"]), None)

    for p in list(nicht_zutreffender_block):
        mark_as_deleted(p)
    if block_art == "OHNE" and ai_table is not None:
        for row in ai_table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    mark_as_deleted(p)
    changes.append((
        "Kap. 3.2: Einsatz von Künstlicher Intelligenz - nicht zutreffender Block (inkl. AI-Standard/High-Tabelle, falls MIT-Block) durchgestrichen/rot markiert",
        f"{'MIT' if block_art == 'OHNE' else 'OHNE'} Künstliche Intelligenz (ganzer Absatzblock)",
        "[durchgestrichen/rot - zum Löschen vorgeschlagen]",
    ))

    kernaussage = "Innerhalb des CS kommt KEINE künstliche Intelligenz zum Einsatz." if block_art == "OHNE" else "Innerhalb des CS kommt künstliche Intelligenz zum Einsatz."
    p = next((p for p in zutreffender_block if p.text.strip() == kernaussage), None)
    if p:
        tag = mark_as_kept(p, bereich="Kap. 3.2")
        changes.append(("Kap. 3.2: KI-Kernaussage trifft laut ki_reifegrad zu", p.text, p.text))
        merke_nummer(tag, "Kap. 3.2: KI-Kernaussage", kernaussage[:50])

    REIFEGRAD_ZEILEN = {
        "I": "KI Reifegrad I: Das KI gestützte CS wird paralle zum Produktionsprozess eingesetzt und keinen Einfluss auf einen GxP relevanten Prozess",
        "II": "KI Reifegrad II: Das KI gestützte CS ist eine konventionelle Anwendung ohne Einsatz von Machine Learning (ML)",
        "III": "KI Reifegrad III: Bei dem KI gestützten CS handelt es ich um Anwendungen die als geschlossenes KI gestützte CS verwendet werden.",
        "IV": "KI Reifegrad IV: Das KI gestützte CS ist autonom mit selbstauslösenden Neutraining – mit menschlichen Eingrif und Kontrolle der Updates",
    }
    paare = ["I", "II"] if block_art == "OHNE" else ["III", "IV"]
    if ki_wert in paare:
        for kandidat in paare:
            p = next((p for p in list(zutreffender_block) if p.text.strip() == REIFEGRAD_ZEILEN[kandidat]), None)
            if not p:
                continue
            if kandidat == ki_wert:
                tag = mark_as_kept(p, bereich="Kap. 3.2")
                changes.append((f"Kap. 3.2: KI Reifegrad {ki_wert} - zutreffende Zeile", p.text, p.text))
                merke_nummer(tag, f"Kap. 3.2: KI Reifegrad {ki_wert}", "zutreffend")
            else:
                delete_paragraph(p, f"Kap. 3.2: KI Reifegrad {kandidat} (ki_reifegrad={ki_wert} trifft nicht zu)")
    elif ki_wert == "N/A":
        NA_ZUSATZZEILEN = [
            "Die „Prohibited AI Practices at Sanofi“ gemäß QU-STD-0000534-AI Systems Risk and Compliance Management ist geprüft und bewertet. Das Ergebnis ist über die Systembewertung dokumentiert.",
            "Gemäß Systembewertung ergibt sich ein KI Reifegrad von I oder II",
            REIFEGRAD_ZEILEN["I"], REIFEGRAD_ZEILEN["II"],
            "Die KI-Kontrollmaßnahmen aus dem QU-OPE-2497575-Risk Profile and Control Framework (Fragen in Abschnitt „5. RISK DOMAIN: Artificial Intelligence“) wurden beantwortet und führten zu dem",
            "Risk Profile Outcome: „N/A“ oder „Non AI“",
        ]
        for zeilentext in NA_ZUSATZZEILEN:
            p = next((p for p in list(zutreffender_block) if p.text.strip() == zeilentext.strip()), None)
            if p:
                delete_paragraph(p, "Kap. 3.2: KI-Reifegrad-Detail/Risk-Profile-Zeile (ki_reifegrad=N/A -> trifft komplett nicht zu)")

    if block_art == "MIT":
        digital_lines = [
            "Die KI-Kontrollmaßnahmen aus dem QU-OPE-2497575-Risk Profile and Control Framework (Fragen in Abschnitt „5. RISK DOMAIN: Artificial Intelligence“) wurden beantwortet und führten zu dem",
            "Risk Profile Outcome: „AI Standard“ oder „AI High“",
            "Gemäß des Risk Profile Outcome ergeben sich aus dem Risk Profile and Control Framework folgende Maßnahmen:",
            "Die Maßnahmen werden in der URS aufgenommen und im CS Validierungsbericht abschließend geprüft.",
        ]
        found_ps = [p for p in doc.paragraphs if p.text.strip() in [t.strip() for t in digital_lines]]
        table_cell_paragraphs = [p for row in ai_table.rows for cell in row.cells for p in cell.paragraphs] if ai_table is not None else []
        if PROJ["digital_beteiligt"] == "ja":
            for p in found_ps:
                delete_paragraph(p, "Kap. 3.2: Risk-Profile-Absatz - digital_beteiligt=ja (ServiceNow-Weg gilt statt dessen)")
            for p in table_cell_paragraphs:
                mark_as_deleted(p)
            if ai_table is not None:
                changes.append(("Kap. 3.2: AI-Standard/High-Tabelle gestrichen (digital_beteiligt=ja, Projekt-DB)", "-", "[durchgestrichen/rot - zum Löschen vorgeschlagen]"))
        else:
            for p in found_ps:
                mark_as_kept(p, bereich="Kap. 3.2")
            if found_ps:
                changes.append(("Kap. 3.2: Risk-Profile-Absatz bleibt erhalten (digital_beteiligt=nein, Projekt-DB)", found_ps[0].text[:60], found_ps[0].text[:60]))
            for p in table_cell_paragraphs:
                mark_as_kept(p)
            if ai_table is not None:
                changes.append(("Kap. 3.2: AI-Standard/High-Tabelle bleibt erhalten (digital_beteiligt=nein, Projekt-DB)", "-", "-"))
    else:
        skipped.append(("Kap. 3.2: 'Wenn Digital beteiligt ist...'-Unterblock (digital_beteiligt)",
            f"ki_reifegrad='{ki_wert}' -> der MIT-Block trifft ohnehin nicht zu und ist schon komplett gestrichen."))

# ============================================================ 12. Kap. 3.4 Designqualifizierung (DQ)
p = find_p(doc, lambda p: p.text.strip().startswith("Es werden / wurden User Requirement Specifications"))
if p:
    ersatz = "wurden" if PROJ["urs_bereits_erstellt"] == "ja" else "werden"
    replace_marker(p, "werden / wurden", ersatz, "Kap. 3.4: URS bereits erstellt? (Projekt-DB)", bereich="Kap. 3.4")

def join_und(items):
    if len(items) <= 1:
        return items[0] if items else ""
    return ", ".join(items[:-1]) + " und " + items[-1]

ausgewaehlte_phasen = [ph for ph in ["DQ", "IQ", "OQ", "PQ", "PPQ"] if PROJ["phasen_geplant"].get(ph)]
ausgewaehlte_testphasen = [ph for ph in ["IQ", "OQ", "PQ", "PPQ"] if PROJ["phasen_geplant"].get(ph)]

p = find_p(doc, lambda p: p.text.strip().startswith("GMP-relevante Punkte werden einer CS Validierung"))
if p:
    replace_marker(p, "(DQ, IQ, OQ, PQ)", "(" + ", ".join(ausgewaehlte_phasen) + ")", "Kap. 3.4: Geplante Phasen (Projekt-DB)", bereich="Kap. 3.4")

p = find_p(doc, lambda p: "Standardtestvorschrift in welcher Projektphase" in p.text)
if p:
    replace_marker(p, "(IQ, OQ, PQ)", "(" + ", ".join(ausgewaehlte_testphasen) + ")", "Kap. 3.4.1: Geplante Testphasen für Standardtestvorschriften (Projekt-DB)", bereich="Kap. 3.4.1")

p = find_p(doc, lambda p: "Abschlussberichte zur IQ, OQ und PQ" in p.text)
if p:
    replace_marker(p, "IQ, OQ und PQ", join_und(ausgewaehlte_testphasen), "Kap. 3.7.4: Geplante Testphasen (UE/Anomalien-Anhang, Projekt-DB)", bereich="Kap. 3.7.4")

if not PROJ["phasen_geplant"].get("PPQ"):
    block = block_between(doc, lambda p: p.style.name == "Heading 3" and p.text.strip() == "Prozess Performance Qualifizierung (PPQ)")
    delete_block(block, "Kap. 3.7.8: Prozess Performance Qualifizierung (PPQ) - komplett gestrichen (phase_ppq_geplant=nein, Projekt-DB)")

p = find_p(doc, lambda p: "zu den Qualifizierungsphasen" in p.text and "(DQ, IQ, OQ, PQ)" in p.text)
if p:
    replace_marker(p, "(DQ, IQ, OQ, PQ)", "(" + ", ".join(ausgewaehlte_phasen) + ")", "Kap. 3.9: Geplante Qualifizierungsphasen im Validierungsbericht (Projekt-DB)", bereich="Kap. 3.9")

p = find_p(doc, lambda p: p.text.strip().startswith("GEP-relevante Punkte können geprüft werden"))
if p:
    if PROJ["gep_pruefung_erforderlich"] == "ja":
        tag = mark_as_kept(p, bereich="Kap. 3.4")
        changes.append(("Kap. 3.4: GEP-Prüfung erforderlich (Projekt-DB) - Absatz bleibt", p.text, p.text))
        merke_nummer(tag, "Kap. 3.4: GEP-Prüfung erforderlich", "ja")
    else:
        delete_paragraph(p, "Kap. 3.4: GEP-Prüfung nicht erforderlich (Projekt-DB) - Absatz gestrichen")

p = find_p(doc, lambda p: "bezogen auf das Projektbezeichnung" in p.text)
if p:
    if PROJ.get("projektbezeichnung"):
        replace_marker(p, "Projektbezeichnung", PROJ["projektbezeichnung"], "Kap. 3.5: Projektbezeichnung (Projekt-DB)", bereich="Kap. 3.5")
    replace_marker(p, "an den Systemen", f"an dem System {SYS['systemname']}", "Kap. 3.5: Systemname statt generisch 'den Systemen' (System-DB)", bereich="Kap. 3.5")

# --- Risikoanalyse: GxP-Kritikalitaet-abhaengiger Absatzblock ---
HOCH = {"Critical", "Major"}
kritikalitaet_hoch = SYS["gxp_kritikalitaet"] in HOCH
RA_HOCH_ZEILEN = [
    "Im Rahmen der Designphase ist eine Risikoanalyse zur Bestimmung der kritischen Prozesschritte auf ihren möglichen Einfluss auf die Produktqualität, die Patientensicherheit und die Datenintegrität zu prüfen und zu erstellen.",
    "Die Risikoanalyse ist mittels einer FMEA (Fehler-Möglichkeits-Einfluss-Analyse oder geeigneten Methoden der Risikoanalyse durchzuführen.",
    "Die Risikoanalyse dient neben dem TM zur Bestimmung des Testumfanges. Die Risikoanalyse dient der Bestimmung der Testtiefe für das jeweilige Risiko. Die Durchführung der Risikoanalyse wird im DQ-Abschlussbericht dokumentiert. Sie stellt sicher, dass mögliche Risiken erkannt und reduziert werden bzw. über Tests, Wartungsaktivitäten oder SOPs abgedeckt werden.",
    "Basierend auf den Ergebnissen der Risikoanalyse und des detaillierten Datenflussschemas ist ein Audit Trail Review Konzept (siehe Kapitel 5) erarbeiten und es ist zu Beschreiben in welcher Form ein regelmäßiger Audit Trail Review erfolgt.",
]
RA_MINOR_ZEILE = "Das System wurde als GxP minor eingestuft. Gemäß QU-SOP-0015430 / QU-SOP-0049866 wird kein RA gefordert und erstellt."
RA_FREIGABE_ZEILE = "Hierzu muss eine freigegebene Risikoanalyse vorliegen, die auf Basis der freigegebenen Spezifikationen durchgeführt wurde."
ra_hoch_tag = None
for text in RA_HOCH_ZEILEN:
    p = find_p(doc, lambda p, t=text: p.text.strip() == t.strip())
    if not p:
        continue
    if kritikalitaet_hoch:
        ra_hoch_tag = mark_as_kept(p, bereich="Kap. 3.4.2") or ra_hoch_tag
    else:
        delete_paragraph(p, f"Kap. 3.4.2: RA-Absatz (major/critical) - GxP-Kritikalität={SYS['gxp_kritikalitaet']}, nicht major/critical -> gestrichen")
if kritikalitaet_hoch:
    changes.append(("Kap. 3.4.2: RA-Absatzblock für major/critical bleibt (GxP-Kritikalität, System-DB)", RA_HOCH_ZEILEN[0][:60], RA_HOCH_ZEILEN[0][:60]))
    if ra_hoch_tag:
        merke_nummer(ra_hoch_tag, "Kap. 3.4.2: RA-Absatzblock major/critical", "zutreffend")
p = find_p(doc, lambda p: p.text.strip() == RA_MINOR_ZEILE)
if p:
    if not kritikalitaet_hoch:
        tag = mark_as_kept(p, bereich="Kap. 3.4.2")
        changes.append(("Kap. 3.4.2: RA-Absatz für minor/N.A. bleibt (GxP-Kritikalität, System-DB)", p.text, p.text))
        merke_nummer(tag, "Kap. 3.4.2: RA-Absatz minor/N.A.", "zutreffend")
    else:
        delete_paragraph(p, f"Kap. 3.4.2: RA-Absatz (minor) - GxP-Kritikalität={SYS['gxp_kritikalitaet']} ist major/critical -> gestrichen")
p = find_p(doc, lambda p: p.text.strip() == RA_FREIGABE_ZEILE)
if p:
    if kritikalitaet_hoch:
        tag = mark_as_kept(p, bereich="Kap. 3.4.3")
        changes.append(("Kap. 3.4.3: 'freigegebene Risikoanalyse muss vorliegen' bleibt (GxP-Kritikalität major/critical)", p.text, p.text))
        merke_nummer(tag, "Kap. 3.4.3: freigegebene Risikoanalyse muss vorliegen", "zutreffend")
    else:
        delete_paragraph(p, f"Kap. 3.4.3: 'freigegebene Risikoanalyse muss vorliegen' - GxP-Kritikalität={SYS['gxp_kritikalitaet']} ist minor/N.A. (keine RA) -> gestrichen")

# ============================================================ 13. Kap. 3.7.1 Testplan (3 Optionen)
TESTPLAN_TEXTE = {
    "separat_freigegeben": "In einem Testplan werden alle durchzuführenden Tests aufgelistet. Der Testplan wird vor Beginn der Testphase separat freigegeben.",
    "als_anhang": "Da bereits eine vollständige Traceability Matrix und eine Risikoanalyse vorliegen, wird der Testplan als Anhang zu diesem Validierungsplan freigegeben.",
    "integriert_im_vp": "Da bereits eine vollständige Traceability Matrix und eine Risikoanalyse vorliegen, ist der Testplan in diesem Validierungsplan integriert. Info: In diesem Fall hier den Testplan einfügen",
}
gewaehlt = PROJ["testplan_art"]
for art, text in TESTPLAN_TEXTE.items():
    p = find_p(doc, lambda p, t=text: p.text.strip().startswith(t.split(" Info:")[0]))
    if not p:
        continue
    if art == gewaehlt:
        tag = mark_as_kept(p, bereich="Kap. 3.7.1")
        changes.append((f"Kap. 3.7.1: Testplan-Variante '{art}' ausgewählt (Projekt-DB)", p.text[:60], p.text[:60]))
        merke_nummer(tag, "Kap. 3.7.1: Testplan-Variante", art)
    else:
        delete_paragraph(p, f"Kap. 3.7.1: Testplan-Variante '{art}' nicht ausgewählt (Projekt-DB)")

# ============================================================ 14. Kap. 3.7.3 Durchführung von Tests (2 Optionen)
TESTDURCHFUEHRUNG_TEXTE = {
    "lieferant_ergebnisse_uebernehmen": "Bei Übernahme von Testergebnissen / Testprotokollen des Lieferanten ist sicherzustellen, dass diese vollständig und richtig sind sowie den GxP-Anforderungen entsprechen. Im V-Plan oder Testplan werden diese Dokumente gelistet.",
    "lieferant_durchfuehrung_sanofi_aufsicht": "Die Durchführung der Tests erfolgt durch den Hersteller/Lieferanten, unter Aufsicht von Sanofi. Nach der Durchführung der Tests werden diese durch Sanofi geprüft und gegengezeichnet, Änderungen ggf. dokumentiert und bewertet.",
}
gewaehlt = PROJ["testdurchfuehrung_art"]
for art, text in TESTDURCHFUEHRUNG_TEXTE.items():
    p = find_p(doc, lambda p, t=text: p.text.strip() == t)
    if not p:
        continue
    if art == gewaehlt:
        tag = mark_as_kept(p, bereich="Kap. 3.7.3")
        changes.append((f"Kap. 3.7.3: Testdurchführung-Variante '{art}' ausgewählt (Projekt-DB)", p.text[:60], p.text[:60]))
        merke_nummer(tag, "Kap. 3.7.3: Testdurchführung-Variante", art)
    else:
        delete_paragraph(p, f"Kap. 3.7.3: Testdurchführung-Variante '{art}' nicht ausgewählt (Projekt-DB)")

# ============================================================ 15. Kap. 3.7.5 Installationsqualifizierung (IQ)
p = find_p(doc, lambda p: "<<Liferant/Sanofi>>" in p.text)
if p:
    replace_marker(p, "<<Liferant/Sanofi>>", PROJ["verantwortlich_iq_durchfuehrung"], "Kap. 3.7.5: Verantwortlich für Durchführung der IQ (Projekt-DB)", bereich="Kap. 3.7.5")

# ============================================================ 16. Tabelle 9: Dokumentenübersicht (Verantwortlichkeiten je Dokument)
VERANTWORTLICH_MAP = {
    "Funktionsspezifikation": PROJ["verantwortlich_funktionsspezifikation"],
    "Hardware Designspezifikation (HDS)": PROJ["verantwortlich_hds"],
    "Software Designspezifikation (SDS)": PROJ["verantwortlich_sds"],
    "Risikoanalyse (RA)": PROJ["verantwortlich_ra"],
    "Traceability-Matrix (TM)": PROJ["verantwortlich_tm"],
    "IQ-Testvorschriften": PROJ["verantwortlich_iq_testvorschriften"],
}
for t in doc.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if header == ["Dokumententyp", "Verantwortlichkeit"]:
        for row in t.rows[1:]:
            dokumenttyp = row.cells[0].text.strip()
            if dokumenttyp in VERANTWORTLICH_MAP and "Firma/Lieferant" in row.cells[1].text:
                set_cell_value(row.cells[1], VERANTWORTLICH_MAP[dokumenttyp], f"Tabelle 9: Dokumentenübersicht - {dokumenttyp} (Projekt-DB)", mark_black=True, bereich="Tabelle 9")
        break

# ============================================================ 17. Kap. 3.9 Systemfreigabe <<System>>
GESCHUETZTE_PARAGRAPHEN = set()
system_freigabe_ps = find_all_p(doc, lambda p: "<<System>>" in p.text)
count_system_marker = 0
for p in system_freigabe_ps:
    replace_marker(p, "<<System>>", SYS["systemname"], f"Kap. 3.9: Systemfreigabe-Absatz (Alternative {count_system_marker+1} von 2 - EINE davon ist vor Freigabe zu löschen)", bereich="Kap. 3.9")
    count_system_marker += 1
if len(system_freigabe_ps) == 2:
    paras = doc.paragraphs
    element_ids = [id(p._element) for p in paras]
    i1 = element_ids.index(id(system_freigabe_ps[0]._element))
    i2 = element_ids.index(id(system_freigabe_ps[1]._element))
    for p in paras[i1 + 1:i2]:
        if p.text.strip() == "Oder":
            mark_as_kept(p)
            GESCHUETZTE_PARAGRAPHEN.add(id(p._element))
            changes.append(("Kap. 3.9: 'Oder' zwischen den beiden Freigabe-Alternativen bleibt stehen (beide Alternativen sind erhalten, Auswahl bleibt Sache des Menschen)", "Oder", "Oder"))

# ============================================================ 18. Tabelle 8: Lieferantenbewertung
for t in doc.tables:
    if t.rows[0].cells[0].text.strip() == "Lieferant / Adresse":
        row = t.rows[1]
        cell0 = row.cells[0]
        old0 = cell0.text
        runs0 = [r for p in cell0.paragraphs for r in p.runs]
        if runs0 and runs0[0].text.strip().startswith("Lieferant A"):
            idx = runs0[0].text.find("Lieferant A")
            tag = nr("Tabelle 8")
            runs0[0].text = f"[{tag}] " + SYS["hersteller"] + runs0[0].text[idx + len("Lieferant A"):]
            runs0[0].font.color.rgb = RGBColor(0, 0, 0)
            runs0[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
            changes.append(("Tabelle 8: Lieferantenbewertung - Lieferantenname", old0.strip(), cell0.text.strip()))
            merke_nummer(tag, "Tabelle 8: Lieferantenbewertung - Lieferantenname", SYS["hersteller"])
        if SYS.get("lieferantennummer"):
            bewertung_cell = row.cells[2]
            for p in bewertung_cell.paragraphs:
                mark_as_deleted(p)
            changes.append(("Tabelle 8: Lieferantenbewertung - 'Ist durchzuführen' gestrichen (QualiPSO-ID liegt vor)", bewertung_cell.text.strip(), "[durchgestrichen/rot - QualiPSO-ID liegt bereits vor]"))
            append_to_cell(bewertung_cell, f"QualiPSO-ID: {SYS['lieferantennummer']}", "Tabelle 8: Lieferantenbewertung - QualiPSO-Nummer in letzter Spalte ergänzt", bereich="Tabelle 8")
        else:
            skipped.append(("Tabelle 8: Lieferantenbewertung - 'Ist durchzuführen'", "Keine QualiPSO-ID (lieferantennummer) hinterlegt -> Bewertung bleibt als ausstehend markiert."))
        break

# ============================================================ 19. V2.0 (CC Nummer)-Zeilen + Versionshistorie (Projekt-DB, FV-Logik)
FV = float(PROJ["folgeversion"])
history_sorted = sorted(VERSIONSHISTORIE, key=lambda e: float(e["version"]))
UEBERGEORDNETE_ZEILE = "V2.0 (CC Nummer): Beschreibung der Änderung, die die Versionierung des V-Plans erforderlich macht."
p60 = find_p(doc, lambda p: p.text.strip() == UEBERGEORDNETE_ZEILE)
if p60:
    if FV <= 1.0 or not history_sorted:
        delete_paragraph(p60, "Kap. 1.1: Versionshistorie (übergeordnete Zeile) - FV=1.0, keine Historie (Projekt-DB)")
    else:
        erster = history_sorted[0]
        neuer_text = f"V{erster['version']} ({erster['cc_nummer']}): {erster['beschreibung']}"
        replace_whole_paragraph(p60, neuer_text, f"Kap. 1.1: Versionshistorie V{erster['version']} (Projekt-DB)", bereich="Kap. 1.1")
        anchor = p60._element
        for eintrag in history_sorted[1:]:
            tag = nr("Kap. 1.1")
            klartext = f"V{eintrag['version']} ({eintrag['cc_nummer']}): {eintrag['beschreibung']}"
            neuer_text = f"[{tag}] {klartext}"
            new_p = doc.add_paragraph(style=p60.style)
            run = new_p.add_run(neuer_text)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            anchor.addnext(new_p._element)
            anchor = new_p._element
            changes.append((f"Kap. 1.1: Versionshistorie V{eintrag['version']} (Projekt-DB, neue Zeile eingefügt, da FV>=3.0)", "-", neuer_text))
            merke_nummer(tag, f"Kap. 1.1: Versionshistorie V{eintrag['version']}", klartext[:50])

V2_CC_RE = re.compile(r"^(V(\d+\.\d+))\s*\((CC[\s-]?Nummer)\)", re.IGNORECASE)
andere_v2_zeilen = find_all_p(doc, lambda p: V2_CC_RE.match(p.text.strip()))
history_by_version = {e["version"]: e for e in VERSIONSHISTORIE}
if FV <= 1.0:
    for p in andere_v2_zeilen:
        delete_paragraph(p, "Kapitelspezifische V2.0(CC Nummer)-Zeile - FV=1.0, keine Historie (Projekt-DB)")
else:
    fehlende_version = []
    for p in andere_v2_zeilen:
        m = V2_CC_RE.match(p.text.strip())
        version = m.group(2)
        eintrag = history_by_version.get(version)
        if eintrag:
            replace_marker(p, f"({m.group(3)})", f"({eintrag['cc_nummer']})", f"Kapitelspezifische V{version}(CC Nummer)-Zeile - CC-Nummer eingesetzt (Projekt-DB), Rest bleibt Freitext", bereich="Kapitelspez. CC-Nr.")
        else:
            fehlende_version.append(version)
    if fehlende_version:
        skipped.append((
            f"Kapitelspezifische 'V{fehlende_version[0]} (CC Nummer): ...'-Zeilen ohne passenden Versionshistorie-Eintrag",
            f"Für Version {', '.join(sorted(set(fehlende_version)))} liegt kein Eintrag in VERSIONSHISTORIE vor -> CC-Nummer bleibt Platzhalter."
        ))
    skipped.append((
        "Die 10 kapitelspezifischen 'V2.0 (CC Nummer): ...'-Zeilen: nur CC-Nummer ersetzt, keine neue V3.0-Zeile",
        f"FV={PROJ['folgeversion']} >= 2.0 -> die '(CC Nummer)'-Platzhalter wurden durch die echte CC-Nummer der "
        "jeweiligen Version ersetzt (Projekt-DB). Die eigentliche Alternative bleibt unangetastet (Freitext)."
    ))

# ============================================================ 20. Tabelle 12: Änderungshistorie
for t in doc.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if len(header) == 3 and header[1] == "Change Control":
        if FV <= 1.0:
            skipped.append(("Tabelle 12: Änderungshistorie", "FV=1.0 (Erstprojekt) - die vorhandene 'Letzte Unterschrift / 1.0'-Zeile ('Neuerstellung') ist bereits korrekt, keine weitere Zeile nötig."))
        else:
            ref_row = t.rows[1]
            for eintrag in history_sorted:
                new_row = t.add_row()
                werte = [f"Freigabedatum: -- / {eintrag['version']}", eintrag["cc_nummer"], eintrag["beschreibung"]]
                labels = [f"Tabelle 12: Änderungshistorie - neue Zeile V{eintrag['version']} (Projekt-DB)", f"Tabelle 12: Änderungshistorie - CC-Nummer V{eintrag['version']}", f"Tabelle 12: Änderungshistorie - Grund V{eintrag['version']}"]
                for ci, (wert, label) in enumerate(zip(werte, labels)):
                    ref_p = ref_row.cells[ci].paragraphs[0]
                    ref_run = next((r for r in ref_p.runs if r.text.strip()), ref_p.runs[0] if ref_p.runs else None)
                    set_cell_value(new_row.cells[ci], wert, label, mark_black=True, bereich="Tabelle 12")
                    neue_p = new_row.cells[ci].paragraphs[0]
                    neue_p.alignment = ref_p.alignment
                    if ref_run is not None:
                        for r in neue_p.runs:
                            r.font.name = ref_run.font.name
                            r.font.size = ref_run.font.size
                            r.font.bold = ref_run.font.bold
        break

# ============================================================ 21. Tabelle 10: Weitere Validierungsdokumente (ECHTE ANKREUZFELDER)
# NEU in v8: die einzige Stelle im CS-VP mit echten Word-Ankreuzfeldern
# (w14:checkbox Content Controls) statt normalem Text. 18 Zeilen x Ja/Nein.
# Aktuell hat NUR die PPQ-Zeile ein passendes Projekt-DB-Feld
# (phase_ppq_geplant) - die anderen 17 haben noch KEIN Feld in der
# Projekt-DB (siehe Modul-Docstring oben) und bleiben deshalb bewusst
# unangekreuzt, statt eine GxP-relevante Ja/Nein-Aussage zu erfinden.
TABELLE10_PRUEFPUNKTE = [
    "Datenflussdiagramm", "Audit Trail Review Konzept", "Berechtigungskonzept",
    "Trainingsplan", "Process Performance Qualification (PPQ)",
    "User Process Monitoring (UPM)", "Datenmigration (DM)",
    "Festlegung Wartung und Monitoring", "Prozedur zur Archivierung der Daten",
    "Back-up & Restore Konzept", "Business Continuity Plan",
    "Incident- und Störungsmanagement", "Änderungs- und Konfigurationsmanagement",
    "Logbuch (Server, CS, Equipment, etc.)", "Lieferantenbewertung",
    "Quality Agreement", "Erstellen von Anweisungen zur Bedienung des CS",
    "Erstellen eines Handbuchs mit detaillierten Anleitungen zur Bedienung des CS",
]
for t in doc.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if header and header[0] == "Prüfpunkte":
        ohne_feld = []
        for i, row in enumerate(t.rows[1:], start=1):
            prueftext = row.cells[0].text.strip()
            if prueftext.startswith("Process Performance Qualification"):
                antwort = "ja" if PROJ["phasen_geplant"].get("PPQ") else "nein"
                set_ja_nein_checkbox(row.cells[1], antwort, "Tabelle 10: Weitere Validierungsdokumente - Prüfpunkt 'Process Performance Qualification (PPQ)' (Projekt-DB: phase_ppq_geplant)", bereich="Tabelle 10")
                if antwort == "nein":
                    append_to_cell(row.cells[2], "nicht erforderlich, da PPQ für dieses Projekt nicht geplant ist (Projekt-DB: phase_ppq_geplant=nein)", "Tabelle 10: Weitere Validierungsdokumente - Begründung PPQ", bereich="Tabelle 10")
            else:
                kurz = TABELLE10_PRUEFPUNKTE[i - 1] if i - 1 < len(TABELLE10_PRUEFPUNKTE) else prueftext[:40]
                ohne_feld.append(kurz)
        skipped.append((
            "Tabelle 10: Weitere Validierungsdokumente - 17 von 18 Prüfpunkten ohne DB-Feld",
            "NEU ENTDECKT bei diesem Durchlauf: nur 'Process Performance Qualification (PPQ)' hat ein passendes "
            "Projekt-DB-Feld (phase_ppq_geplant). Für die folgenden 17 Prüfpunkte gibt es noch KEIN Feld in der "
            "Projekt-DB, die Ankreuzfelder bleiben deshalb bewusst unangetastet (☐/☐, keine erfundene GxP-Aussage): "
            + ", ".join(ohne_feld) + ". Wären eigene Ja/Nein-Felder in seed_field_definitions_projekt.sql (analog "
            "vmp_erforderlich/gep_pruefung_erforderlich), sollte das aufgenommen werden - noch nicht umgesetzt."
        ))
        break

# ============================================================ 22. Grauen Text (#A6A6A6) im gesamten Dokument durchstreichen/rot markieren
def run_color(r):
    return r.font.color.rgb if r.font.color and r.font.color.type else None

def strip_grey(paragraphs_iterable):
    markiert = 0
    for p in list(paragraphs_iterable):
        if id(p._element) in GESCHUETZTE_PARAGRAPHEN:
            continue
        grey_runs = [r for r in p.runs if run_color(r) == RGBColor(0xA6, 0xA6, 0xA6) and r.text]
        if not grey_runs:
            continue
        for r in grey_runs:
            r.font.strike = True
            r.font.color.rgb = DELETE_COLOR
        markiert += 1
    return markiert

grau_markiert = strip_grey(doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            grau_markiert += strip_grey(cell.paragraphs)
changes.append(("Grauer Text (#A6A6A6) im gesamten Dokument durchgestrichen/rot markiert", f"{grau_markiert} Absätze/Zellen betroffen", "[durchgestrichen/rot - zum Löschen vorgeschlagen]"))

# ============================================================ 23. Aufräumen: doppelte Leerzeilen nach dem Löschen
def _ist_leer(p):
    return p.text.strip() == ""

def _ist_komplett_durchgestrichen(p):
    runs_mit_text = [r for r in p.runs if r.text]
    return bool(runs_mit_text) and all(r.font.strike for r in runs_mit_text)

geloescht_count = 0
survivors = [p for p in doc.paragraphs if not _ist_komplett_durchgestrichen(p)]
i = 0
while i < len(survivors) - 1:
    if _ist_leer(survivors[i]) and _ist_leer(survivors[i + 1]):
        survivors[i + 1]._element.getparent().remove(survivors[i + 1]._element)
        del survivors[i + 1]
        geloescht_count += 1
    else:
        i += 1

i = 0
while i < len(survivors) - 1:
    if _ist_leer(survivors[i]) and survivors[i + 1].style.name.startswith("Heading"):
        survivors[i]._element.getparent().remove(survivors[i]._element)
        del survivors[i]
        geloescht_count += 1
    else:
        i += 1

if geloescht_count:
    changes.append(("Aufräumen: überzählige Leerzeilen entfernt (nach simuliertem Löschen keine 2 Returns in Folge, keine Leerzeile direkt vor einer Überschrift)", f"{geloescht_count} leere Absätze entfernt", "[entfernt]"))

for p in doc.paragraphs:
    if _ist_leer(p):
        for r in p.runs:
            r.font.color.rgb = None
            r.font.highlight_color = None
            r.font.strike = None

doc.save(OUT)

# ============================================================ Zuordnungstabelle (NEU in v8) - Nr. -> Ort/Feld -> Wert
with open(ZUORDNUNGSDATEI, "w", encoding="utf-8") as f:
    f.write("# Zuordnungstabelle CS-VP Demo v8 (Nr. im Dokument -> Feld/Ort -> Wert)\n\n")
    f.write("Jede Nummer in eckigen Klammern im Dokument (z.B. `[Tabelle 1-1]`) entspricht\n")
    f.write("genau einer Zeile hier - reine Abstimmungshilfe, vor echter Nutzung wieder zu entfernen.\n\n")
    f.write("| Nr. | Ort/Feld (Label) | Eingesetzter Wert |\n")
    f.write("|---|---|---|\n")
    for tag, label, wert in NUMMERIERUNG:
        f.write(f"| `{tag}` | {label} | {wert} |\n")

print("Fertig:", OUT)
print("Zuordnungstabelle:", ZUORDNUNGSDATEI)
print(f"\n{len(NUMMERIERUNG)} nummerierte Stellen, {len(changes)} automatische Änderungen insgesamt:")
for label, old, new in changes:
    print(f"  [{label}]")
    print(f"     vorher:  {str(old)[:80]!r}")
    print(f"     nachher: {str(new)[:80]!r}")
print(f"\n{len(skipped)} bewusst NICHT (voll) angefasste Stellen (Begründung):")
for label, reason in skipped:
    print(f"  [{label}]\n     {reason}\n")

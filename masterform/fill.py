"""Füllen von reinen Text-Templates in Word-Dokumenten (.docx).

Ein "reines Text-Template" ist hier ein normales Word-Dokument, in dem
Platzhalter der Form ``{{feld_name}}`` als ganz normaler Text stehen -
im Gegensatz zu Dokumenten mit echten Formularfeldern/Content Controls
(Steuerelementen), die eine andere Verarbeitung benötigen.

Die Hauptschwierigkeit beim Ersetzen von Text in .docx-Dateien mit
``python-docx`` ist, dass Word einen einzelnen zusammenhängenden Text oft in
mehrere ``run``-Objekte aufteilt (z.B. wegen Rechtschreibprüfung oder weil
beim Tippen die Formatierung minimal wechselt). Ein naives
``run.text.replace(...)`` je Run findet einen Platzhalter dann nicht, wenn
er über mehrere Runs verteilt ist. Die Funktionen hier arbeiten daher auf
dem zusammengesetzten Paragraph-Text und schreiben das Ergebnis
Run-genau zurück, sodass die vorhandene Formatierung möglichst erhalten
bleibt.
"""

from __future__ import annotations

import re
from typing import Any, Iterable, Mapping

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# Platzhalter-Syntax: {{feld}}, {{ feld }}, {{feld_name_123}}
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([A-Za-z0-9_.\-]+)\s*\}\}")


def find_placeholders(path: str) -> set[str]:
    """Liefert alle im Dokument vorkommenden Platzhalter-Namen (ohne Werte)."""
    doc = Document(path)
    names: set[str] = set()
    for paragraph in _iter_all_paragraphs(doc):
        names.update(m.group(1) for m in PLACEHOLDER_PATTERN.finditer(paragraph.text))
    return names


def fill_template(template_path: str, output_path: str, data: Mapping[str, Any]) -> None:
    """Öffnet ein Template, ersetzt alle bekannten Platzhalter und speichert es.

    Args:
        template_path: Pfad zur Word-Vorlage (.docx) mit Platzhaltern wie
            ``{{name}}``.
        output_path: Zielpfad für die ausgefüllte Datei.
        data: Zuordnung Platzhalter-Name -> Wert. Werte werden mit ``str()``
            in Text umgewandelt. Platzhalter, für die kein Wert übergeben
            wurde, bleiben unverändert im Text stehen.
    """
    doc = Document(template_path)
    fill_document(doc, data)
    doc.save(output_path)


def fill_document(doc: DocumentObject, data: Mapping[str, Any]) -> None:
    """Ersetzt Platzhalter direkt in einem bereits geöffneten ``Document``."""
    for paragraph in _iter_all_paragraphs(doc):
        _replace_in_paragraph(paragraph, data)


def _iter_all_paragraphs(doc: DocumentObject) -> Iterable[Paragraph]:
    """Iteriert über alle Absätze: Fließtext, Tabellen, Kopf- und Fußzeilen."""
    yield from doc.paragraphs
    yield from _iter_table_paragraphs(doc.tables)

    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if part is None:
                continue
            yield from part.paragraphs
            yield from _iter_table_paragraphs(part.tables)


def _iter_table_paragraphs(tables: Iterable[Table]) -> Iterable[Paragraph]:
    """Iteriert rekursiv über alle Absätze in Tabellen (inkl. verschachtelter)."""
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    yield from cell.paragraphs
    yield from _iter_table_paragraphs(cell.tables)


def _replace_in_paragraph(paragraph: Paragraph, data: Mapping[str, Any]) -> bool:
    """Ersetzt alle bekannten Platzhalter in einem Absatz, Run-genau.

    Verarbeitet Treffer von hinten nach vorne, damit sich durch
    unterschiedlich lange Ersatzwerte die Positionen der noch nicht
    bearbeiteten Treffer nicht verschieben.
    """
    text = paragraph.text
    matches = list(PLACEHOLDER_PATTERN.finditer(text))
    if not matches:
        return False

    replaced_any = False
    for match in reversed(matches):
        key = match.group(1)
        if key not in data:
            continue  # unbekannte Platzhalter unangetastet lassen
        value = str(data[key])
        if _replace_span_in_paragraph(paragraph, match.start(), match.end(), value):
            replaced_any = True
    return replaced_any


def _run_offsets(paragraph: Paragraph) -> list[tuple[int, int]]:
    offsets = []
    pos = 0
    for run in paragraph.runs:
        offsets.append((pos, pos + len(run.text)))
        pos += len(run.text)
    return offsets


def _replace_span_in_paragraph(paragraph: Paragraph, start: int, end: int, value: str) -> bool:
    """Ersetzt den Textbereich [start, end) des Absatzes durch ``value``.

    Die Formatierung des ersten betroffenen Runs bleibt für den neuen Wert
    erhalten. Runs, die vollständig im ersetzten Bereich liegen, werden
    geleert statt gelöscht (einfacher & robust genug für Template-Zwecke).
    """
    runs = paragraph.runs
    offsets = _run_offsets(paragraph)
    overlapping = [i for i, (s, e) in enumerate(offsets) if e > start and s < end]
    if not overlapping:
        return False

    first_idx, last_idx = overlapping[0], overlapping[-1]
    first_run = runs[first_idx]
    first_s, _first_e = offsets[first_idx]
    prefix = first_run.text[: start - first_s]

    if first_idx == last_idx:
        suffix = first_run.text[end - first_s :]
        first_run.text = prefix + value + suffix
        return True

    last_run = runs[last_idx]
    last_s, _last_e = offsets[last_idx]
    suffix = last_run.text[end - last_s :]

    first_run.text = prefix + value
    for i in overlapping[1:-1]:
        runs[i].text = ""
    last_run.text = suffix
    return True
